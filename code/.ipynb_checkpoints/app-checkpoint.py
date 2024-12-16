import streamlit as st
import joblib
import pandas as pd
from PIL import Image
from docx import Document  # For DOCX files
#import fitz  # For PDF files
import zipfile
import io
import mammoth  # For DOCX text extraction
from sklearn.feature_extraction.text import TfidfVectorizer
from pythainlp.tokenize import word_tokenize
from pythainlp.corpus.common import thai_words, thai_stopwords
from pythainlp.util import dict_trie
from sklearn.preprocessing import LabelEncoder


# Load the model and its components
objects = joblib.load('../model/model.pkl')
tfidf = objects['tfidf']
model = objects['model']
label_encoder = objects['label_encoder']

# Sample texts for testing
POE_SAMPLE = """Once upon a midnight dreary, while I pondered, weak and weary, 
Over many a quaint and curious volume of forgotten lore..."""

AUSTEN_SAMPLE = """It a truth universally acknowledged, that a single man in possession of good fortune,
must be in want of a wife..."""

# Sidebar for file upload
st.sidebar.header("File Upload")
uploaded_files = st.sidebar.file_uploader("Upload files", 
                                         type=["csv", "xlsx", "docx", "pdf", "zip"], 
                                         accept_multiple_files=True)



# Display the title and description
st.title("Logistics Paper Classification")
st.subheader("Do you want to know what category your research paper belongs to?")

with st.expander("How it works"):
    st.write("This app uses machine learning to categorize your research paper.")
    st.write("Tips for best results:")
    st.write("- Enter at least 10 characters")
    st.write("- The more text, the better the predictive power")
    st.write("- Try sample text to see how it works")

def extract_abstract_from_docx(uploaded_file):
    keyword = 'บทคัดย่อ'
    stop_phrases = ['บทนำ','ที่มาและความสำคัญ','1ที่มา','คำสำคัญ','นิยามศัพท์เฉพาะ','ทบทวนวรรณกรรม']
    captured_lines = []     
    try:
        # Use BytesIO to read the uploaded file as a byte stream
        doc = Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Split the extracted text into lines
        lines = text.split('\n')
        contents = [text.strip() for text in lines if text != '']  # All contents
        
        # Flag to start capturing text after finding the first matching line
        capture_text = False

        for line in lines:
            if not capture_text:
                # Check if line contains any of the keywords
                if line.strip() == keyword:
                    capture_text = True  # Start capturing from this line onward
                elif keyword in line:
                    capture_text = True
                    captured_lines.append(line)  # Start capturing from this line onward
                 
            else:
                # Stop capturing if the line contains a stop phrase
                if any(stop_phrase in line for stop_phrase in stop_phrases):
                    break
                captured_lines.append(line)
        
        # Remove unwanted lines
        captured_lines = [item for item in captured_lines if item != keyword and item.strip() != '']
        abstract = '\n'.join(captured_lines)
        st.header("Abstract: ")
        
        # Split the abstract into chunks of 200 characters or less and join them with line breaks
        max_line_length = 150
        formatted_abstract = "\n".join([abstract[i:i + max_line_length] for i in range(0, len(abstract), max_line_length)])
        
        # Display the formatted abstract in one go
        st.text(formatted_abstract)
        
        # Preprocess and classify text
        preprocess_and_predict(text)
        
    except Exception as e:
        st.error(f"Error extracting abstract: {e}")
        
# Extract abstract from PDF
def extract_abstract_from_pdf(uploaded_file):
    try:
        pdf_text = []
        with fitz.open(uploaded_file) as doc:
            for page in doc:
                pdf_text.append(page.get_text())
        text = "\n".join(pdf_text)
        # Example abstract extraction
        start = text.find("บทคัดย่อ")  # Adjust if necessary
        if start != -1:
            abstract = text[start:]
            st.write("Abstract Extracted:")
            st.text(abstract)
            preprocess_and_predict(abstract)
    except Exception as e:
        st.error(f"Error extracting abstract: {e}")

def preprocess_and_predict(text):
    # Add custom word to keep from 5 sample abstract
    added_words = ['การนำเข้า', 'ฐานนิยม', 'คราฟท์', 'แนวทาง', 'ผู้ส่งมอบ', 'โซ่อุปทาน', 'ปัจจัยรอง', 
                   'การส่งมอบ', 'รถขนส่ง', 'นำไปใช้งาน', 'อย่างถูกต้อง', 'การขับรถ', 'ที่เกี่ยวข้อง', 
                   'ในการปฏิบัติงาน', 'พนักงานขับรถ', 'สิ่งสำคัญ', 'ขั้นตอน', 'ที่ชัดเจน', 'การไหล', 'ยอดขาย', 
                   'การจัดทำ', 'คราฟท์เบียร์', 'ฝึกสหกิจ', 'อย่างก้าวกระโดด','การจัดซื้อจัดหา','กระบวนการ',
                   'แบบประเมิน','เก็บข้อมูล','อย่างชัดเจน','การดำเนินการ','การส่งเสริม','ถังดับเพลิง','แนวทาง']
    
    # Merge custom words with Thai dictionary words
    custom_words = set(thai_words()).union(added_words)
    custom_trie = dict_trie(custom_words)  # Create a trie from the custom dictionary
 
    tokens = word_tokenize(text, custom_dict= custom_trie, engine="newmm")
    text =  " ".join([token for token in tokens if token not in thai_stopwords()])
    
    # Example of tokenization or text preprocessing
    text_tfidf = tfidf.transform([text])  # Transform your input text
    
    prediction = model.predict(text_tfidf.toarray(), verbose = 0)  # Convert sparse matrix to dense array
    
    # Decode class labels
    all_classes = label_encoder.inverse_transform(range(prediction.shape[1]))
    
    # Combine classes and probabilities into a list of tuples
    class_probs = list(zip(all_classes, prediction[0]))
    
    # Sort class_probs by the second element (probability) in descending order
    sorted_class_probs = sorted(class_probs, key=lambda x: x[1], reverse=True)
    
    # Filter out classes with probabilities less than 40%
    filtered_class_probs = [cat for cat in sorted_class_probs if cat[1] * 100 > 40]
    
    # If no class has a probability greater than 40%, display only the top class
    if not filtered_class_probs:
        filtered_class_probs = [sorted_class_probs[0]]
    
    # Display the filtered classes and show a progress bar
    
    for cat in filtered_class_probs:
        col1, col2 = st.columns([3, 1])  # Adjust column ratio to suit your needs
    
        with col1:
            st.write(f"Class: {cat[0]}")
            st.write(f'Probability: {cat[1] * 100:.2f}%')
    
        with col2:
            # Show progress bar for the probability
            st.progress(float(cat[1]))  # The value should be between 0 and 1
    st.markdown('---')

# Main page for results
if uploaded_files:
    st.sidebar.success(f"{len(uploaded_files)} file(s) uploaded successfully!")
    
    # Iterate over the uploaded files and process them
    for uploaded_file in uploaded_files:
        st.write(f"File Name: {uploaded_file.name}")
        
        if uploaded_file.name.endswith(".docx"):
            # Process DOCX files
            doc = Document(uploaded_file)
            full_text = [para.text for para in doc.paragraphs]
            extract_abstract_from_docx(uploaded_file)  # Extract abstract and classify text

        elif uploaded_file.name.endswith(".pdf"):
            # Process PDF files (implementation needed for PDF extraction)
            extract_abstract_from_pdf(uploaded_file)

        elif uploaded_file.name.endswith(".csv"):
            # Process CSV files
            df = pd.read_csv(uploaded_file)
            st.dataframe(df)

        elif uploaded_file.name.endswith(".xlsx"):
            # Process Excel files
            df = pd.read_excel(uploaded_file)
            st.dataframe(df)

        else:
            st.warning("Unsupported file format")

